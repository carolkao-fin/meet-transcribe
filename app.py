"""
MeetTranscribe — Streamlit 版（使用 Groq，完全免費）
上傳音檔（最大 100 MB）→ Groq Whisper 轉錄 → 指定發言者 → Groq Llama 分析
"""

import base64
import io
import json
import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
from streamlit_javascript import st_javascript

# ── 頁面設定 ───────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="MeetTranscribe",
    page_icon="🎙",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.hero {
    background: linear-gradient(135deg, #1BA8A8 0%, #138A8A 100%);
    color: white; border-radius: 14px;
    padding: 1.4rem 2rem; margin-bottom: 1.5rem;
}
.hero h1 { color: white; margin: 0; font-size: 1.9rem; font-weight: 800; }
.hero p  { color: rgba(255,255,255,.85); margin: .3rem 0 0; font-size: .95rem; }
.res-header {
    background: linear-gradient(135deg, #1BA8A8 0%, #138A8A 100%);
    color: white; border-radius: 12px;
    padding: 1.4rem 1.6rem; margin-bottom: 1.2rem;
}
.res-header h2 { color: white; margin: 0 0 .4rem; font-size: 1.4rem; }
.res-header .meta { font-size: .88rem; opacity: .88; }
.pbadge {
    background: rgba(255,255,255,.22);
    padding: .15rem .65rem; border-radius: 20px;
    font-size: .8rem; font-weight: 700;
    display: inline-block; margin: .15rem .2rem 0 0;
}
.sec-title {
    color: #1BA8A8; font-size: .82rem; font-weight: 800;
    text-transform: uppercase; letter-spacing: .6px;
    border-bottom: 2px solid #e0f5f5;
    padding-bottom: .35rem; margin: 1.4rem 0 .7rem;
}
.sum-h   { font-weight: 700; font-size: .97rem; margin: .9rem 0 .25rem; color: #1f2937; }
.sum-ov  { color: #6b7280; font-size: .9rem; margin-bottom: .4rem; line-height: 1.62; }
.bullet  { font-size: .9rem; padding-left: 1.1rem; position: relative; margin: .22rem 0; line-height: 1.56; }
.bullet::before { content:"·"; position:absolute; left:0; color:#1BA8A8; font-weight:700; font-size:1.1rem; }
.ttag {
    background: #e0f5f5; color: #138A8A;
    padding: .22rem .75rem; border-radius: 20px;
    font-size: .85rem; font-weight: 600;
    display: inline-block; margin: .2rem .2rem 0 0;
}
.ai-title { font-weight: 700; font-size: .95rem; margin: .9rem 0 .2rem; }
.ai-desc  { color: #6b7280; font-size: .85rem; margin-bottom: .45rem; }
.ai-item  { font-size: .9rem; padding: .18rem 0 .18rem 1.1rem; position: relative; line-height:1.55; }
.ai-item::before { content:"·"; position:absolute; left:0; color:#1BA8A8; font-weight:700; }
.ai-who   { font-weight: 700; color: #138A8A; }
.tr-row   { display:flex; gap:.75rem; padding:.28rem 0; font-size:.9rem; line-height:1.57; }
.tr-spk   { font-weight:700; min-width:90px; flex-shrink:0; color:#1BA8A8; }
.hist-card {
    border: 1.5px solid #e5e7eb; border-radius: 10px;
    padding: .75rem 1rem; margin-bottom: .5rem; background: white;
}
.hist-card.active { border-color: #1BA8A8; background: #e0f5f5; }
.hist-title { font-weight: 700; font-size: .95rem; color: #1f2937; }
.hist-meta  { font-size: .78rem; color: #9ca3af; margin-top: .2rem; }
</style>
""", unsafe_allow_html=True)

# ── Session state ──────────────────────────────────────────────────────────────
for k, v in {
    "transcript":        [],
    "analysis":          None,
    "meeting_info":      {},
    "history":           [],
    "hist_idx":          None,
    "current_record_id": None,
    "_last_uploaded":    "",
    "_history_loaded":   False,
    "_last_txt_upload":  "",
    "docx_bytes":        None,
    "docx_name":         "會議紀錄.docx",
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

if "meeting_title_key" not in st.session_state:
    st.session_state["meeting_title_key"] = f"會議記錄 {datetime.now():%Y-%m-%d}"

# 若上次 rerun 有待套用的檔名，在 sidebar 渲染前先更新 widget key
if st.session_state.get("_pending_title"):
    st.session_state["meeting_title_key"] = st.session_state.pop("_pending_title")

# ── 從 localStorage 載入歷史記錄（每個 session 只做一次）────────────────────────
if not st.session_state["_history_loaded"]:
    _raw = st_javascript('localStorage.getItem("meetTranscribeHistory")')
    if _raw and _raw != 0 and _raw not in ("null", "undefined"):
        try:
            _loaded = json.loads(_raw)
            if isinstance(_loaded, list) and _loaded:
                st.session_state.history = _loaded
        except Exception:
            pass
    if _raw != 0:          # 0 表示 JS 還沒執行完，等下次 rerun 再標記
        st.session_state["_history_loaded"] = True


def _persist_history() -> None:
    """將目前歷史記錄同步至瀏覽器 localStorage（音訊檔案不存入，避免超過大小限制）。"""
    slim = [{k: v for k, v in r.items() if k not in ("audio_bytes",)} for r in st.session_state.history]
    data = json.dumps(slim, ensure_ascii=False)
    components.html(
        f'<script>localStorage.setItem("meetTranscribeHistory",{json.dumps(data)});</script>',
        height=0,
    )

# ── 常數 ───────────────────────────────────────────────────────────────────────
LANG_MAP      = {"自動偵測": None, "中文 (zh)": "zh", "英文 (en)": "en"}
WHISPER_MAX   = 24 * 1_048_576          # 24 MB
WHISPER_MODEL = "whisper-large-v3"

# 切片健康檢查：小於 MIN_CHUNK_BYTES 視為 ffmpeg 產出的空檔，
# 短於 MIN_CHUNK_SEC 的尾段直接捨棄（Whisper 無法處理零長度音訊）
MIN_CHUNK_BYTES = 2048
MIN_CHUNK_SEC   = 1.0

# Groq 轉錄端點能直接解析的容器。注意不含 raw .aac（ADTS），
# 副檔名不在此清單的檔案一律先用 ffmpeg 轉成 mp3 再送。
GROQ_AUDIO_EXTS = {
    ".flac", ".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".ogg", ".opus", ".wav", ".webm",
}
CHAT_MODEL    = "llama-3.3-70b-versatile"

# Groq 免費方案 TPM 上限 ~12,000 tokens；
# prompt 固定部分約 500 tokens，max_tokens=2048，故逐字稿最多保留 ~9,000 tokens。
# 中文約 1.5 chars/token → 安全字元上限取 12,000 chars
MAX_TRANSCRIPT_CHARS = 12_000

# ── 會議紀錄 .docx 樣式（對應「會議行動轉譯官」skill 的 SeaSalt.AI 規範）──────
DOCX_TEAL    = "00897B"              # 主色
DOCX_BODY    = "212121"              # 正文色
DOCX_FONT    = "Arial"               # 拉丁字元
DOCX_FONT_EA = "Microsoft JhengHei"  # 中日韓字元（Arial 沒有中文字符，會退回 Word 預設）
DOCX_TITLE_PT = 22
DOCX_H1_PT    = 18
DOCX_H2_PT    = 14
DOCX_META_PT  = 11
DOCX_BD_PT    = 12                   # body

# H2 自動編號用的中文數字（一、二、三…）
_CJK_NUM = "一二三四五六七八九十"
# 已經帶編號的標題不再重複加（「1. 」「一、」「(1)」…）
_NUM_PREFIX_RE = re.compile(r"^\s*(?:\d+[.\)、]|[（(]\d+[）)]|[" + _CJK_NUM + r"]+[、.])\s*")

# 預設五段結構；Transcript 由程式填入完整逐字稿，不交給 LLM 產生
DEFAULT_SECTIONS = ["Meeting Summary", "Topics", "Action Items", "Notes"]

# ── Helpers ────────────────────────────────────────────────────────────────────
def secs_hms(s: float) -> str:
    s = int(s)
    return f"{s//3600:02d}:{(s%3600)//60:02d}:{s%60:02d}"


def _friendly_error(e: Exception) -> str:
    """將 Groq API 錯誤轉成友善的中文訊息。"""
    msg = str(e)
    # 429 rate limit
    if "429" in msg or "rate_limit_exceeded" in msg:
        wait = re.search(r"try again in\s+([\d]+m[\d.]+s|[\d.]+s|[\d]+m)", msg)
        wait_str = f"請等待 **{wait.group(1)}** 後再試。" if wait else "請稍後再試。"
        if "ASPH" in msg or "seconds of audio" in msg:
            return f"⏱ Groq 免費方案每小時音訊轉錄量已達上限（7,200 秒）。{wait_str}"
        if "TPM" in msg or "tokens per minute" in msg:
            return f"⏱ Groq 免費方案每分鐘 Token 用量已達上限。{wait_str}"
        return f"⏱ Groq API 請求次數已達上限。{wait_str}"
    # 413 too large
    if "413" in msg:
        return "📏 逐字稿過長，已超過 Groq 模型的 Token 上限，請縮短錄音後再試。"
    # 400 檔案無法解碼
    if "could not process file" in msg or "valid media file" in msg:
        return (
            "🎧 Groq 無法解析這個音訊檔，可能原因：\n\n"
            "- 副檔名與實際格式不符（例如 m4a 被改名成 .mp3）\n"
            "- 格式不在 Groq 支援清單內：flac / mp3 / mp4 / mpeg / mpga / m4a / "
            "ogg / opus / wav / webm（**不含 raw .aac**）\n"
            "- 檔案損毀或內容為空\n\n"
            "建議先確認檔案能正常播放，或轉存成 mp3 / wav 後再上傳。"
        )
    return msg


def _ffmpeg_ok() -> bool:
    try:
        r = subprocess.run(["ffmpeg", "-version"], capture_output=True, timeout=5)
        return r.returncode == 0
    except Exception:
        return False


def _audio_duration(path: str) -> float:
    try:
        r = subprocess.run(
            ["ffprobe", "-v", "quiet", "-print_format", "json",
             "-show_streams", "-show_format", path],
            capture_output=True, text=True, timeout=30,
        )
        info = json.loads(r.stdout)
        dur = next(
            (float(s["duration"]) for s in info.get("streams", []) if "duration" in s), 0.0
        )
        # webm / ogg 等容器常缺少 stream duration，退回 format duration
        if not dur:
            dur = float(info.get("format", {}).get("duration") or 0.0)
        return dur
    except Exception:
        return 0.0


def _split_audio(src: str, chunk_min: int = 8) -> list | None:
    """切成數段 mp3。ffmpeg / ffprobe 不可用時回傳 None；
    切片本身失敗則 raise，並帶上 ffmpeg 的 stderr 方便診斷。"""
    if not _ffmpeg_ok():
        return None
    dur = _audio_duration(src)
    if not dur:
        return None

    # 用去掉原副檔名的 base，避免產生 xxx.m4a_c0.mp3 這種雙副檔名
    base = str(Path(src).with_suffix(""))
    chunks, step, t, idx = [], chunk_min * 60, 0.0, 0
    while t < dur:
        length = min(step, dur - t)
        if length < MIN_CHUNK_SEC:      # 捨棄零長度尾段，否則 ffmpeg 產出空檔
            break
        out = f"{base}_c{idx}.mp3"
        # -ss 放在 -i 之前：input seeking，不必每段都從頭解碼
        r = subprocess.run(
            ["ffmpeg", "-y", "-ss", str(t), "-i", src, "-t", str(length),
             "-ac", "1", "-ar", "16000", "-b:a", "48k", out],
            capture_output=True, text=True, timeout=300,
        )
        size = Path(out).stat().st_size if Path(out).exists() else 0
        if r.returncode != 0 or size < MIN_CHUNK_BYTES:
            for done, _ in chunks:      # 清掉已產生的切片，別留垃圾
                try:
                    os.unlink(done)
                except Exception:
                    pass
            try:
                os.unlink(out)
            except Exception:
                pass
            tail = " / ".join((r.stderr or "").strip().splitlines()[-3:]) or "無"
            raise RuntimeError(
                f"ffmpeg 分割第 {idx+1} 段失敗"
                f"（returncode={r.returncode}，輸出 {size} bytes）。\n\nffmpeg 訊息：{tail}"
            )
        chunks.append((out, length))
        t += length
        idx += 1
    return chunks or None


def _transcode_to_mp3(src: str) -> str:
    """把 Groq 不支援的容器（例如 raw .aac）轉成 mp3。失敗則 raise。"""
    out = f"{Path(src).with_suffix('')}_conv.mp3"
    r = subprocess.run(
        ["ffmpeg", "-y", "-i", src, "-ac", "1", "-ar", "16000", "-b:a", "48k", out],
        capture_output=True, text=True, timeout=600,
    )
    size = Path(out).stat().st_size if Path(out).exists() else 0
    if r.returncode != 0 or size < MIN_CHUNK_BYTES:
        try:
            os.unlink(out)
        except Exception:
            pass
        tail = " / ".join((r.stderr or "").strip().splitlines()[-3:]) or "無"
        raise RuntimeError(
            f"ffmpeg 轉檔失敗（returncode={r.returncode}，輸出 {size} bytes）。"
            f"\n\nffmpeg 訊息：{tail}"
        )
    return out


def _call_whisper(client, path: str, lc: str | None) -> list[dict]:
    size = os.path.getsize(path)
    if size < MIN_CHUNK_BYTES:
        raise ValueError(
            f"要送出的音訊只有 {size} bytes，內容為空或已損毀，已中止上傳。"
        )
    with open(path, "rb") as fh:
        r = client.audio.transcriptions.create(
            model=WHISPER_MODEL,
            file=fh,
            response_format="verbose_json",
            timestamp_granularities=["segment"],
            language=lc,
        )
    out = []
    for s in getattr(r, "segments", None) or []:
        start = float(s["start"] if isinstance(s, dict) else s.start)
        text  = (s["text"]  if isinstance(s, dict) else getattr(s, "text", "") or "").strip()
        if text:
            out.append({"start": start, "text": text})
    if not out and getattr(r, "text", None):
        out.append({"start": 0.0, "text": r.text.strip()})
    return out


def transcribe_audio(data: bytes, filename: str, groq_key: str, lc: str | None) -> list[dict]:
    from groq import Groq
    client = Groq(api_key=groq_key)
    suffix = Path(filename).suffix.lower() or ".mp3"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    tmp_files = [tmp_path]
    try:
        # Groq 不認識的容器（raw .aac 等）先轉成 mp3，否則會回 400 invalid media file
        if suffix not in GROQ_AUDIO_EXTS:
            if not _ffmpeg_ok():
                raise ValueError(
                    f"Groq 不支援 {suffix} 格式，且找不到 ffmpeg 無法自動轉檔。\n"
                    "請先自行轉存成 mp3 或 wav 再上傳。"
                )
            with st.spinner(f"{suffix} 非 Groq 支援格式，正在轉成 mp3…"):
                tmp_path = _transcode_to_mp3(tmp_path)
            tmp_files.append(tmp_path)

        size = os.path.getsize(tmp_path)
        if size <= WHISPER_MAX:
            return _call_whisper(client, tmp_path, lc)

        chunks = _split_audio(tmp_path)
        if not chunks:
            raise ValueError(
                f"檔案 {size/1_048_576:.1f} MB 超過 24 MB，需要 ffmpeg 自動分割，"
                "但找不到 ffmpeg／ffprobe，或無法讀取這個檔案的音訊長度。\n"
                "請確認檔案格式正常，或先壓縮成 24 MB 以下再上傳。"
            )
        segs, offset = [], 0.0
        pbar = st.progress(0, text="分割並轉錄中…")
        try:
            for i, (cp, dur) in enumerate(chunks):
                pbar.progress((i + 1) / len(chunks), text=f"轉錄第 {i+1}/{len(chunks)} 段…")
                for s in _call_whisper(client, cp, lc):
                    segs.append({"start": s["start"] + offset, "text": s["text"]})
                offset += dur
        finally:
            pbar.empty()
            for cp, _ in chunks:        # 中途失敗也要清掉暫存切片
                try:
                    os.unlink(cp)
                except Exception:
                    pass
        return segs
    finally:
        for p in tmp_files:
            try:
                os.unlink(p)
            except Exception:
                pass


def analyze_with_groq(transcript: list, meeting_info: dict, groq_key: str) -> dict:
    from groq import Groq
    client = Groq(api_key=groq_key)

    txt      = "\n".join(f"{e['speaker']}: {e['text']}" for e in transcript)
    has_zh   = any("\u4e00" <= c <= "\u9fff" for e in transcript for c in e["text"])
    out_lang = "繁體中文" if has_zh else "English"

    # ── 超長逐字稿截斷（保留前 2/3、後 1/3，避免 TPM 超限）──────────────────
    truncated = False
    if len(txt) > MAX_TRANSCRIPT_CHARS:
        keep_head = int(MAX_TRANSCRIPT_CHARS * 0.67)
        keep_tail = MAX_TRANSCRIPT_CHARS - keep_head
        txt = txt[:keep_head] + f"\n\n… [逐字稿過長，中間部分已略去] …\n\n" + txt[-keep_tail:]
        truncated = True

    prompt = f"""你是一個專業的會議記錄分析師。請分析以下會議逐字稿，並以 {out_lang} 輸出結構化的分析結果。{"（注意：逐字稿因過長已截取首尾，請根據現有內容盡力分析。）" if truncated else ""}

會議資訊：
- 標題：{meeting_info.get("title", "未命名會議")}
- 日期：{meeting_info.get("date", "")}
- 參與者：{", ".join(meeting_info.get("participants", []))}

逐字稿：
{txt}

請根據上下文修正專業術語、公司名稱、產品名稱、人名等用詞。

只回傳一個合法的 JSON 物件，不要包含 markdown 或其他文字：

{{
  "summary": [
    {{
      "title": "### 小節標題",
      "overview": "這一小節的概述段落",
      "bullets": ["- speaker_X 說明了...", "- 雙方討論了...", "- 最終確認..."]
    }}
  ],
  "topics": ["主題標籤1", "主題標籤2"],
  "action_items": [
    {{
      "group_title": "### 行動任務群組標題",
      "description": "這組任務的說明",
      "items": [{{"assignee": "負責人", "task": "具體任務"}}]
    }}
  ],
  "corrected_transcript": [{{"speaker": "speaker_1", "text": "修正後文字"}}]
}}

要求：summary 2–4 小節、每節 3–6 要點並標明發言者；action_items 指明負責人；topics 2–4 個標籤。"""

    resp = client.chat.completions.create(
        model=CHAT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
        temperature=0.3,
    )
    text = resp.choices[0].message.content.strip()
    if "```json" in text:
        text = text.split("```json")[1].split("```")[0].strip()
    elif text.startswith("```"):
        text = text.split("```")[1].split("```")[0].strip()
    return json.loads(text)


def save_to_history(transcript, analysis, meeting_info, record_id=None,
                    audio_bytes=None, audio_filename=None) -> str:
    """新增或更新歷史記錄，回傳 record id。"""
    if record_id:
        for r in st.session_state.history:
            if r["id"] == record_id:
                r["transcript"]   = transcript
                r["analysis"]     = analysis
                r["meeting_info"] = meeting_info
                r["participants"] = meeting_info.get("participants", [])
                if audio_bytes is not None:
                    r["audio_bytes"]    = audio_bytes
                    r["audio_filename"] = audio_filename or "recording"
                return record_id

    record = {
        "id":             datetime.now().strftime("%Y%m%d_%H%M%S"),
        "title":          meeting_info.get("title", "未命名會議"),
        "date":           meeting_info.get("date", ""),
        "participants":   meeting_info.get("participants", []),
        "transcript":     transcript,
        "analysis":       analysis,
        "meeting_info":   meeting_info,
        "audio_bytes":    audio_bytes,
        "audio_filename": audio_filename or "recording",
    }
    st.session_state.history.insert(0, record)
    return record["id"]


def plain_text(data: dict, info: dict, transcript: list) -> str:
    lines = [
        info.get("title", "會議記錄"),
        f"日期：{info.get('date', '')}",
        f"參與者：{', '.join(f'[{p}]' for p in info.get('participants', []))}",
        "", "Meeting Summary", "=" * 40,
    ]
    for sec in data.get("summary", []):
        lines += [sec.get("title", ""), sec.get("overview", "")]
        lines += sec.get("bullets", [])
        lines.append("")
    lines += ["Topics", ", ".join(data.get("topics", [])), "", "Action Items", "=" * 40]
    for grp in data.get("action_items", []):
        lines += [grp.get("group_title", ""), grp.get("description", "")]
        lines += [f"- {i['assignee']} {i['task']}" for i in grp.get("items", [])]
        lines.append("")
    lines += ["Transcript", "=" * 40]
    src = data.get("corrected_transcript") or transcript
    lines += [f"{e['speaker']}: {e['text']}" for e in src]
    return "\n".join(lines)


# ── 逐字稿 .txt → 會議紀錄 .docx ────────────────────────────────────────────────
def _decode_text(raw: bytes) -> str:
    """逐字稿多半是 UTF-8，但 Windows 記事本另存可能是 CP950（Big5）。"""
    for enc in ("utf-8-sig", "utf-8", "cp950"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


# 「speaker_1: 內容」「王小明：內容」，行首可有 [00:01:23] 之類的時間戳
_SPK_RE = re.compile(
    r"^\s*(?:[\[(]?\d{1,2}:\d{2}(?::\d{2})?[\])]?\s*)?([^：:\n]{1,24})[：:]\s*(.+)$"
)
_TIME_ONLY_RE = re.compile(r"^\s*[\[(]?\d{1,2}:\d{2}(?::\d{2})?[\])]?\s*")


def parse_transcript_txt(text: str) -> list[dict]:
    """把純文字逐字稿拆成 [{"speaker","text"}]。

    - 「發言者：內容」開新的一段，行首時間戳會被去掉
    - 緊接的斷行視為同一段的折行，併入前一段
    - 空行結束該段；空行後沒有發言者標記的內容自成一段（speaker 留空），
      不沿用上一位發言者，避免把話算到錯的人身上
    """
    entries: list[dict] = []
    after_blank = True
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            after_blank = True
            continue
        m = _SPK_RE.match(line)
        if m:
            entries.append({"speaker": m.group(1).strip(), "text": m.group(2).strip()})
        elif entries and not after_blank:
            entries[-1]["text"] += " " + _TIME_ONLY_RE.sub("", line)
        else:
            entries.append({"speaker": "", "text": _TIME_ONLY_RE.sub("", line)})
        after_blank = False
    return entries


_SENT_END  = "。！？!?；;…"
_PUNCT_END = _SENT_END + "，,、：:"


def _join_text(a: str, b: str) -> str:
    """接兩段語音片段。

    Whisper 中文輸出常整份沒有標點，直接黏起來會變成幾百字讀不動的字串。
    這裡在片段之間留一個空白標示停頓——不自行補逗號句號，
    因為斷句會改變語意，那是發言者的話，不該由工具代為決定。
    """
    if not a:
        return b
    return a + b if a[-1] in _PUNCT_END else f"{a} {b}"


def merge_segments(entries: list[dict], max_chars: int = 250) -> list[dict]:
    """把連續同一發言者的短句合併成段落。

    Whisper 逐字稿是一句一行（例如本專案「下載逐字稿」的輸出），
    直接寫進 Word 會變成幾百個各掛一次發言者名稱的單句段落，
    送進 LLM 時重複的前綴也會吃掉大量 token。

    段落長度以 max_chars 為目標，但**只在原始行的邊界斷開**——那是 Whisper 的
    語音停頓點，從那裡斷不會切在句子中間。有句末標點時優先在標點處收段。
    Whisper 中文輸出常整份沒有標點，所以不能只靠標點切。
    續段的 `cont=True`，輸出時不重複印發言者名稱。
    """
    turns: list[dict] = []
    for e in entries:
        if turns and turns[-1]["speaker"] == e["speaker"]:
            turns[-1]["parts"].append(e["text"])
        else:
            turns.append({"speaker": e["speaker"], "parts": [e["text"]]})

    out: list[dict] = []
    for turn in turns:
        para, first = "", True

        def flush(nonlocal_para: str, is_first: bool):
            out.append({"speaker": turn["speaker"], "text": nonlocal_para.strip(),
                        "cont": not is_first})

        for part in turn["parts"]:
            too_long   = para and len(para) + len(part) > max_chars
            sentence_e = para and para[-1] in _SENT_END and len(para) >= max_chars * 0.6
            if too_long or sentence_e:
                flush(para, first)
                para, first = part, False
            else:
                para = _join_text(para, part)
        if para.strip():
            flush(para, first)
    return out


def read_template_text(upload) -> str:
    """讀取使用者上傳的格式模板（.txt / .md / .docx），回傳純文字結構描述。"""
    # UploadedFile 跨 rerun 是同一個物件，用 getvalue() 才不會受讀取位置影響
    raw = upload.getvalue()
    if Path(upload.name).suffix.lower() == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return _decode_text(raw)


def generate_minutes(entries: list[dict], info: dict, groq_key: str,
                     template_text: str | None = None) -> dict:
    """用 Groq Llama 依「忠實還原」原則產生會議紀錄各段落。

    回傳 {"sections": [{"heading": str, "blocks": [...]}]}；
    blocks 的 type 為 h2 / p / bullets。Transcript 不在此產生。
    """
    from groq import Groq
    client = Groq(api_key=groq_key)

    txt = "\n".join(
        f"{e['speaker']}: {e['text']}" if e["speaker"] and not e.get("cont") else e["text"]
        for e in entries
    )
    truncated = False
    if len(txt) > MAX_TRANSCRIPT_CHARS:
        keep_head = int(MAX_TRANSCRIPT_CHARS * 0.67)
        keep_tail = MAX_TRANSCRIPT_CHARS - keep_head
        txt = txt[:keep_head] + "\n\n… [逐字稿過長，中間部分已略去] …\n\n" + txt[-keep_tail:]
        truncated = True

    if template_text:
        fmt_rule = (
            "請嚴格依照下列使用者提供的格式模板決定段落標題與欄位順序，"
            "不得自行新增模板以外的欄位；模板有欄位但逐字稿沒有對應資訊時，填入「——」。\n\n"
            f"格式模板：\n{template_text[:3000]}"
        )
    else:
        fmt_rule = """請採用以下四段結構（Transcript 由系統另外附上完整內容，你不需要輸出）：

1. Meeting Summary
   - 先用一段 p 總述全場：會議性質、雙方是誰、涵蓋哪些議題
   - 再依議題分節，每節：一個 h2 小標題 → 一句 p 說明本節討論什麼 → bullets 條列各方論點
   - 條列每句都要指明是誰說的（例如「公司代表說明…」「訪視方確認…」），
     並盡量保留數字、公司名、代碼等具體資訊
   - 議題請切細一點，寧可多分幾節，不要把不同主題混在同一節

2. Topics
   - 一段 p，以頓號分隔的關鍵詞列表，涵蓋全場核心名詞（人名、公司、產品、制度、代碼）

3. Action Items
   - 依任務群組分節：h2 小標題 → 一句 p 說明 → bullets 條列
   - 條列格式「【負責人】具體任務（期限：…）」；逐字稿沒提到期限就寫
     「（期限：未於逐字稿中提及）」
   - 只列逐字稿中真的要求或承諾要做的事；沒有就整段寫「——」

4. Notes
   - 依性質分節（h2 小標題 + bullets），需要哪幾類就寫哪幾類：
     發言者標記與會議基本資訊（誰是誰、哪些欄位逐字稿沒提到）、
     專有名詞統一說明（同一個詞的多種辨識寫法統一成什麼）、
     語音辨識錯字修正彙整（「錯字 → 正確」逐條列出）、
     詞彙／語意待確認項目、數字與代碼記錄、雜訊與非公務段落說明"""

    prompt = f"""你是一名「會議行動轉譯官」，負責把會議逐字稿轉為正式會議紀錄。

最高原則是**忠實還原**：只呈現逐字稿中實際出現的資訊。
禁止新增未提及的決議或行動事項、禁止靠邏輯推斷補充內容、禁止更改任何人的立場或結論。
逐字稿無對應資訊的欄位，一律填「——」，不得捏造。
{"（注意：逐字稿因過長已截取首尾，請就現有內容整理，不要臆測被略去的部分。）" if truncated else ""}

會議資訊：
- 標題：{info.get("title", "未命名會議")}
- 日期：{info.get("date", "")}

{fmt_rule}

逐字稿：
{txt}

只回傳一個合法的 JSON 物件，最外層必須正好是 {{"sections": [...]}}，
不要包含 markdown、說明文字，也不要用其他鍵名包住 sections：

{{
  "sections": [
    {{
      "heading": "段落標題",
      "blocks": [
        {{"type": "h2", "text": "小節標題"}},
        {{"type": "p", "text": "段落文字"}},
        {{"type": "bullets", "items": ["條列項目1", "條列項目2"]}}
      ]
    }}
  ]
}}

輸出語言：繁體中文（台灣慣用語）；若逐字稿為英文則以英文輸出。"""

    last_raw = ""
    for attempt in range(2):
        msgs = [{"role": "user", "content": prompt}]
        if attempt:        # 第一次形狀不對，把它自己的輸出貼回去要求改正
            msgs += [
                {"role": "assistant", "content": last_raw[:2000]},
                {"role": "user", "content":
                    "格式不對。請只回傳 {\"sections\": [{\"heading\": ..., \"blocks\": [...]}]}，"
                    "最外層鍵名必須是 sections，內容不變。"},
            ]
        resp = _chat_json(client, msgs)
        last_raw = (resp.choices[0].message.content or "").strip()
        if getattr(resp.choices[0], "finish_reason", "") == "length":
            raise ValueError(
                "AI 回應在寫到一半時達到長度上限，內容不完整。"
                "請縮短逐字稿，或取消勾選 AI 整理只做格式轉檔。"
            )
        try:
            sections = _normalize_sections(json.loads(_extract_json(last_raw)))
            if sections:
                return {"sections": sections}
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    raise ValueError(
        "AI 連續兩次都沒有回傳可用的結構，這通常是模型當下不穩定，重試一次多半就好。"
        "若持續失敗，可取消勾選 AI 整理只做格式轉檔。\n\n"
        f"模型實際回傳的開頭：\n{last_raw[:400] or '（空白）'}"
    )


def _chat_json(client, messages, max_tokens: int = 2048):
    """呼叫 Groq 並要求 JSON 輸出；舊模型不支援 response_format 時自動退回一般模式。"""
    kwargs = dict(model=CHAT_MODEL, messages=messages, max_tokens=max_tokens, temperature=0.2)
    try:
        return client.chat.completions.create(response_format={"type": "json_object"}, **kwargs)
    except Exception as e:
        if "response_format" not in str(e) and "json_object" not in str(e):
            raise
        return client.chat.completions.create(**kwargs)


def _extract_json(text: str) -> str:
    """從模型輸出中取出 JSON：去掉 code fence，再抓第一個完整的 {...} 或 [...]。"""
    if "```" in text:
        parts = text.split("```")
        for part in parts[1:]:
            body = part[4:] if part.lower().startswith("json") else part
            if body.strip().startswith(("{", "[")):
                text = body
                break
    text = text.strip()
    start = min((i for i in (text.find("{"), text.find("[")) if i >= 0), default=-1)
    if start < 0:
        return text
    opener = text[start]
    closer = "}" if opener == "{" else "]"
    depth, in_str, esc = 0, False, False
    for i, ch in enumerate(text[start:], start):
        if in_str:
            if esc:
                esc = False
            elif ch == "\\":
                esc = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
        elif ch == opener:
            depth += 1
        elif ch == closer:
            depth -= 1
            if depth == 0:
                return text[start:i + 1]
    return text[start:]


# 這些鍵是會議metadata，不是段落，轉換時要跳過
_META_KEYS = {"title", "date", "meeting_title", "meeting_date", "meeting", "標題", "日期", "會議標題"}


def _item_text(item) -> str:
    """條列項目可能是字串，也可能是 {assignee, task, due} 之類的物件。"""
    if isinstance(item, str):
        return item.lstrip("-• ").strip()
    if isinstance(item, dict):
        who  = item.get("assignee") or item.get("owner") or item.get("負責人") or ""
        task = item.get("task") or item.get("text") or item.get("item") or item.get("事項") or ""
        due  = item.get("due") or item.get("deadline") or item.get("期限") or ""
        if who or task:
            s = f"【{who}】{task}" if who else str(task)
            return f"{s}（期限：{due}）" if due else s
        return "；".join(f"{k}：{v}" for k, v in item.items() if v)
    return str(item)


def _norm_blocks(raw) -> list[dict]:
    """把各種可能的 blocks 形狀正規化成 h2 / p / bullets。"""
    if raw is None or raw == "":
        return []
    if isinstance(raw, str):
        return [{"type": "p", "text": raw.strip()}]
    if isinstance(raw, dict):
        if any(k in raw for k in ("type", "text", "items", "bullets")):
            raw = [raw]
        else:                       # {小標題: 內容} 的對應表
            out = []
            for k, v in raw.items():
                out.append({"type": "h2", "text": str(k)})
                out += _norm_blocks(v)
            return out
    if not isinstance(raw, list):
        return [{"type": "p", "text": str(raw)}]
    if raw and all(isinstance(b, str) for b in raw):
        return [{"type": "bullets", "items": [_item_text(b) for b in raw]}]

    out: list[dict] = []
    for b in raw:
        if isinstance(b, str):
            out.append({"type": "p", "text": b.strip()})
            continue
        if not isinstance(b, dict):
            out.append({"type": "p", "text": str(b)})
            continue
        btype = str(b.get("type", "")).lower()
        items = b.get("items") or b.get("bullets")
        if btype in ("h2", "h3", "heading", "subheading") or (b.get("heading") and not b.get("text")):
            out.append({"type": "h2", "text": str(b.get("text") or b.get("heading")).lstrip("# ")})
        elif isinstance(items, list):
            if b.get("text"):
                out.append({"type": "p", "text": str(b["text"])})
            out.append({"type": "bullets", "items": [_item_text(i) for i in items]})
        elif "text" in b:
            out.append({"type": "p", "text": str(b["text"])})
        else:
            out += _norm_blocks({k: v for k, v in b.items() if k != "type"})
    return out


def _normalize_sections(data) -> list[dict]:
    """接受模型各種常見回傳形狀，統一成 [{"heading", "blocks"}]。

    Llama 不見得會照著 sections 這個鍵名回，實際看過的形狀包括：
    直接回陣列、包在別的鍵名底下、或用 {段落標題: 內容} 的對應表。
    """
    if isinstance(data, list):
        raw_sections = data
    elif isinstance(data, dict):
        if isinstance(data.get("sections"), list):
            raw_sections = data["sections"]
        else:
            nested = next(
                (v for v in data.values()
                 if isinstance(v, list) and v and isinstance(v[0], dict)
                 and ({"heading", "blocks", "title", "content"} & set(v[0]))),
                None,
            )
            if nested is not None:
                raw_sections = nested
            elif len(data) == 1 and isinstance(next(iter(data.values())), dict):
                return _normalize_sections(next(iter(data.values())))
            else:   # {段落標題: 內容} 的對應表
                raw_sections = [
                    {"heading": k, "blocks": v}
                    for k, v in data.items()
                    if str(k).lower() not in _META_KEYS
                ]
    else:
        return []

    out: list[dict] = []
    for s in raw_sections:
        if not isinstance(s, dict):
            continue
        heading = (s.get("heading") or s.get("title") or s.get("name")
                   or s.get("section") or s.get("段落") or "")
        body = None
        for key in ("blocks", "content", "body", "sections", "內容"):
            if key in s:
                body = s[key]
                break
        if body is None:
            body = {k: v for k, v in s.items()
                    if k not in ("heading", "title", "name", "section", "段落")}
        blocks = _norm_blocks(body)
        if heading or blocks:
            out.append({"heading": str(heading).lstrip("# ").strip() or "——", "blocks": blocks})
    return out


# ── 逐字稿整理（補標點、標記錯字、分辨發言者角色）──────────────────────────────
POLISH_CHUNK_CHARS = 1_200      # 每批送出的逐字稿字元數
POLISH_TPM_BUDGET  = 10_000     # 保守值，Groq 免費方案上限 12,000 TPM
POLISH_MIN_RATIO   = 0.6        # 整理後短於原文這個比例 → 視為模型在摘要，改用原文


def _detect_roles(client, sample: str, title: str) -> list[str]:
    """先看開頭一段，決定整份逐字稿要用的角色名稱，避免各批叫法不一致。"""
    prompt = f"""以下是一場會議「{title}」逐字稿的開頭。請判斷這場對話有哪些角色
（例如「訪視方」「公司代表」「主持人」；若能從內容判斷單位或公司名稱，請用具體名稱）。

只回傳 JSON：{{"roles": ["角色1", "角色2"]}}，最多 4 個，不確定就少列。

逐字稿開頭：
{sample[:1500]}"""
    try:
        resp = _chat_json(client, [{"role": "user", "content": prompt}], max_tokens=200)
        roles = json.loads(_extract_json(resp.choices[0].message.content or ""))["roles"]
        return [str(r) for r in roles if str(r).strip()][:4]
    except Exception:
        return []


def polish_transcript(entries: list[dict], groq_key: str, title: str = "",
                      on_progress=None) -> tuple[list[dict], list[str]]:
    """整理逐字稿：補標點、修正明顯辨識錯字並標記【原文：X】、依語意標出發言者角色。

    分批送出並控制速率（Groq 免費方案 12,000 TPM）。任何一批失敗或模型把內容
    縮短了，該批一律保留原文——寧可沒整理，也不能讓逐字稿少內容。

    回傳 (整理後段落, 警告訊息)。
    """
    import time
    from groq import Groq

    client   = Groq(api_key=groq_key)
    warnings: list[str] = []

    # 分批：以字元數切，不切斷單一段落
    chunks: list[list[int]] = []
    cur, cur_len = [], 0
    for i, e in enumerate(entries):
        if cur and cur_len + len(e["text"]) > POLISH_CHUNK_CHARS:
            chunks.append(cur)
            cur, cur_len = [], 0
        cur.append(i)
        cur_len += len(e["text"])
    if cur:
        chunks.append(cur)

    roles = _detect_roles(client, "\n".join(e["text"] for e in entries[:20]), title)
    role_rule = (
        f"發言者角色只能從這幾個裡面選：{'、'.join(roles)}；"
        "若某段無法判斷，填「與會者（姓名未辨識）」。"
        if roles else "請依語意判斷發言者角色；無法判斷時填「與會者（姓名未辨識）」。"
    )

    out = [dict(e) for e in entries]
    used: list[tuple[float, int]] = []      # (時間, token 數)，用來控速

    for n, idx_list in enumerate(chunks, 1):
        if on_progress:
            on_progress(n, len(chunks))

        # 速率控制：讓最近 60 秒的用量維持在預算內
        now = time.time()
        used = [(t, k) for t, k in used if now - t < 60]
        if sum(k for _, k in used) > POLISH_TPM_BUDGET and used:
            wait = 60 - (now - used[0][0])
            if wait > 0:
                if on_progress:
                    on_progress(n, len(chunks), f"（等待 Groq 速率限制，約 {int(wait)} 秒）")
                time.sleep(min(wait, 60))
                used = []

        src = [{"i": i, "text": entries[i]["text"]} for i in idx_list]
        prompt = f"""你是會議逐字稿整理員。以下是語音辨識產生的逐字稿片段（會議：{title}）。

請對每一段做這四件事，其他一律不要動：
1. 補上標點符號（原檔幾乎沒有標點）
2. 修正明顯的語音辨識錯字，並在該處標記【原文：辨識結果】
3. 刪除「呃」「那個」這類純贅詞與重複語句
4. 判斷這段是誰說的。{role_rule}

嚴格禁止：
- 禁止摘要、濃縮或刪除任何實質內容，每一段整理後的長度應與原文相當
- 禁止新增原文沒有的資訊
- 無法判讀的破碎語句：保留原文並在句末標記（語意待確認），不要自行編造

只回傳 JSON，i 必須與輸入一致：
{{"paragraphs": [{{"i": 0, "speaker": "角色", "text": "整理後的文字"}}]}}

逐字稿片段：
{json.dumps(src, ensure_ascii=False)}"""

        try:
            resp = _chat_json(client, [{"role": "user", "content": prompt}], max_tokens=2048)
            usage = getattr(resp, "usage", None)
            used.append((time.time(), getattr(usage, "total_tokens", 3000) if usage else 3000))
            if getattr(resp.choices[0], "finish_reason", "") == "length":
                warnings.append(f"第 {n} 批因長度上限被截斷，該批保留原文。")
                continue
            data = json.loads(_extract_json(resp.choices[0].message.content or ""))
            paras = data.get("paragraphs") or data.get("segments") or []
            got = {int(p["i"]): p for p in paras if isinstance(p, dict) and "i" in p}
            for i in idx_list:
                p = got.get(i)
                if not p:
                    continue
                text = str(p.get("text", "")).strip()
                # 內容縮水就不採用——逐字稿不得被摘要
                if len(text) < len(entries[i]["text"]) * POLISH_MIN_RATIO:
                    continue
                out[i]["text"] = text
                role = str(p.get("speaker", "")).strip()
                if role:
                    base = entries[i]["speaker"]
                    out[i]["speaker"] = f"{base}（{role}）" if base else role
                    out[i]["cont"] = False
        except Exception as e:
            msg = str(e)
            if "429" in msg or "rate_limit" in msg:
                warnings.append(
                    f"第 {n}/{len(chunks)} 批起遇到 Groq 速率限制，其餘段落保留原文。"
                    "稍後再試可完成整理。"
                )
                break
            warnings.append(f"第 {n} 批整理失敗（{msg[:80]}），該批保留原文。")

    return out, warnings


def _style_run(run, size: int = DOCX_BD_PT, color: str = DOCX_BODY, bold: bool = False):
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = DOCX_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    # 中日韓字元要另外指定 eastAsia：Arial 沒有中文字符，
    # 不指定的話中文會退回 Word 預設字型（通常是新細明體）
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT_EA)
    return run


def _cjk_num(n: int) -> str:
    """1→一、10→十、11→十一（會議紀錄的小節編號用不到 100 以上）。"""
    if n <= 10:
        return _CJK_NUM[n - 1]
    if n < 20:
        return "十" + _CJK_NUM[n - 11]
    tens, ones = divmod(n, 10)
    return _CJK_NUM[tens - 1] + "十" + (_CJK_NUM[ones - 1] if ones else "")


def _add_bottom_border(paragraph):
    """H1 標題下方的分隔線。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), DOCX_TEAL)
    pbdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pbdr)


def build_minutes_docx(title: str, date: str, sections: list[dict],
                       entries: list[dict], include_transcript: bool = True,
                       source_name: str = "", transcript_note: str = "") -> bytes:
    """依 skill 的樣式規範產生 .docx，回傳 bytes。

    版面規格對齊既有的人工整理成果（訪談/必佳訪視_會議紀錄.docx）：
    標題 22pt、H1 18pt 加底線、H2 14pt、正文 12pt，
    H1 自動編號「1. 」、H2 自動編號「一、」。
    """
    from docx import Document
    from docx.shared import Cm, Inches, Pt

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)          # A4
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(DOCX_BD_PT)

    def h1(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        _style_run(p.add_run(text), DOCX_H1_PT, DOCX_TEAL, bold=True)
        _add_bottom_border(p)

    def h2(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        _style_run(p.add_run(text), DOCX_H2_PT, DOCX_TEAL, bold=True)

    def body(text: str, size: int = DOCX_BD_PT, color: str = DOCX_BODY):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _style_run(p.add_run(text), size, color)

    def bullet(text: str):
        # 用 numbering 樣式產生項目符號，不手動插入「‧」
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _style_run(p.add_run(text))

    # ── 標題區 ────────────────────────────────────────────────────────────
    head = (title or "會議").strip()
    if not head.endswith("會議紀錄"):
        head = f"{head} 會議紀錄"
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    _style_run(title_p.add_run(head), DOCX_TITLE_PT, DOCX_TEAL, bold=True)
    _add_bottom_border(title_p)

    meta = f"會議日期：{date}" if date else "會議日期：（未於逐字稿中提及）"
    meta += "　｜　時間／地點／主席：（未於逐字稿中提及）"
    if source_name:
        meta += f"　｜　來源檔案：{source_name}"
    body(meta, DOCX_META_PT)

    # ── 各段落（H1 依序編號，H2 於每個 H1 內重新編號）────────────────────
    for n, s in enumerate(sections, 1):
        raw_head = _NUM_PREFIX_RE.sub("", str(s.get("heading", "")).lstrip("# ").strip())
        h1(f"{n}. {raw_head or '——'}")
        sub = 0
        for blk in s.get("blocks", []):
            btype = blk.get("type", "p")
            if btype == "h2":
                sub += 1
                t = _NUM_PREFIX_RE.sub("", str(blk.get("text", "")).lstrip("# ").strip())
                h2(f"{_cjk_num(sub)}、{t}")
            elif btype == "bullets":
                for item in blk.get("items", []):
                    bullet(str(item).lstrip("-• ").strip())
            else:
                body(str(blk.get("text", "")).strip())

    # ── Transcript：完整逐字稿，不經 LLM 摘要 ─────────────────────────────
    if include_transcript:
        h1(f"{len(sections) + 1}. Transcript")
        if transcript_note:
            body(transcript_note, DOCX_META_PT)
        for e in entries:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            # 續段（同一位發言者被拆成多段）不重複印名字
            if e["speaker"] and not e.get("cont"):
                _style_run(p.add_run(f"{e['speaker']}："), DOCX_BD_PT, DOCX_TEAL, bold=True)
            _style_run(p.add_run(e["text"]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def minutes_filename(title: str, date_str: str = "") -> str:
    """{會議主題}_{YYYYMMDD}_會議紀錄.docx；無日期則省略日期段。"""
    safe = re.sub(r'[\\/:*?"<>|]', "_", (title or "會議紀錄").strip()).replace(" ", "_")
    return f"{safe}_{date_str}_會議紀錄.docx" if date_str else f"{safe}_會議紀錄.docx"


def render_results(data: dict, info: dict, transcript: list, key_prefix: str = "main",
                   audio_bytes: bytes | None = None, audio_filename: str = "recording") -> None:
    participants = info.get("participants", [])
    badges = " ".join(f'<span class="pbadge">[{p}]</span>' for p in participants)
    st.markdown(
        f'<div class="res-header"><h2>{info.get("title","會議記錄")}</h2>'
        f'<div class="meta">🕐 {info.get("date","")} &nbsp;·&nbsp; 參與者：{badges}</div></div>',
        unsafe_allow_html=True,
    )
    if audio_bytes:
        st.markdown('<div class="sec-title">🎵 原始錄音</div>', unsafe_allow_html=True)
        st.audio(audio_bytes)
        st.download_button("⬇ 下載原始錄音", audio_bytes, audio_filename,
                           key=f"{key_prefix}_dl_audio")
    if data.get("summary"):
        st.markdown('<div class="sec-title">📋 Meeting Summary</div>', unsafe_allow_html=True)
        for sec in data["summary"]:
            st.markdown(f'<div class="sum-h">### {sec.get("title","").replace("### ","")}</div>', unsafe_allow_html=True)
            if sec.get("overview"):
                st.markdown(f'<div class="sum-ov">{sec["overview"]}</div>', unsafe_allow_html=True)
            for b in sec.get("bullets", []):
                st.markdown(f'<div class="bullet">{b.lstrip("- ")}</div>', unsafe_allow_html=True)
    if data.get("topics"):
        st.markdown('<div class="sec-title">🏷 Topics</div>', unsafe_allow_html=True)
        st.markdown(" ".join(f'<span class="ttag">{t}</span>' for t in data["topics"]), unsafe_allow_html=True)
    if data.get("action_items"):
        st.markdown('<div class="sec-title">✅ Action Items</div>', unsafe_allow_html=True)
        for grp in data["action_items"]:
            st.markdown(f'<div class="ai-title">### {grp.get("group_title","").replace("### ","")}</div>', unsafe_allow_html=True)
            if grp.get("description"):
                st.markdown(f'<div class="ai-desc">{grp["description"]}</div>', unsafe_allow_html=True)
            for item in grp.get("items", []):
                st.markdown(
                    f'<div class="ai-item"><span class="ai-who">{item["assignee"]}</span> {item["task"]}</div>',
                    unsafe_allow_html=True,
                )
    st.markdown('<div class="sec-title">💬 Transcript</div>', unsafe_allow_html=True)
    src = data.get("corrected_transcript") or transcript
    st.markdown(
        "".join(f'<div class="tr-row"><span class="tr-spk">{e["speaker"]}:</span><span>{e["text"]}</span></div>' for e in src),
        unsafe_allow_html=True,
    )
    st.divider()
    fname = info.get("title", "meeting").replace(" ", "_")
    c1, c2, _ = st.columns([2, 2, 4])
    with c1:
        st.download_button("⬇ 下載會議記錄 (.txt)", plain_text(data, info, transcript), f"{fname}.txt",
                           type="primary", key=f"{key_prefix}_dl_txt")
    with c2:
        st.download_button("⬇ 下載 JSON（可重新載入）",
                           json.dumps({"transcript": transcript, "analysis": data, "meeting_info": info},
                                      ensure_ascii=False, indent=2), f"{fname}.json",
                           key=f"{key_prefix}_dl_json")


# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🎙 MeetTranscribe")
    st.caption("智能會議轉錄 & AI 分析")
    st.divider()

    with st.expander("🔑 API 金鑰", expanded=True):
        groq_key = st.text_input(
            "Groq API Key（免費）",
            type="password", placeholder="gsk_...",
            help="免費申請：console.groq.com，一個 key 同時用於轉錄和分析",
        )
        st.caption("📌 [免費取得 Groq Key](https://console.groq.com) — 用 Google 帳號即可註冊")

    with st.expander("📋 會議資訊", expanded=True):
        meeting_title = st.text_input("標題", key="meeting_title_key")
        language      = st.selectbox("語言", ["自動偵測", "中文 (zh)", "英文 (en)"])

    with st.expander("👤 發言者", expanded=True):
        n_sp = st.number_input("人數", 1, 6, 2, step=1)
        speaker_names = [
            st.text_input(f"發言者 {i+1}", value=f"speaker_{i+1}", key=f"spn{i}")
            for i in range(int(n_sp))
        ]

    st.divider()
    if st.button("🗑 清除重來", use_container_width=True):
        st.session_state.transcript        = []
        st.session_state.analysis          = None
        st.session_state.meeting_info      = {}
        st.session_state.current_record_id = None
        st.session_state["_last_uploaded"] = ""
        st.session_state["_pending_title"] = f"會議記錄 {datetime.now():%Y-%m-%d}"
        st.rerun()

# ── Hero ───────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
  <h1>🎙 MeetTranscribe</h1>
  <p>上傳音訊 · Groq Whisper 免費轉錄 · 指定發言者 · Llama AI 免費分析 · 一個 Key 搞定所有功能</p>
</div>
""", unsafe_allow_html=True)

# ── Tabs ───────────────────────────────────────────────────────────────────────
tab_up, tab_rec, tab_docx, tab_hist = st.tabs(
    ["📁 上傳音檔", "🎤 即時錄音", "📄 逐字稿轉 Word", "📚 歷史記錄"]
)

# ── Tab 1: Upload ──────────────────────────────────────────────────────────────
with tab_up:
    st.markdown("支援格式：**mp3 · wav · m4a · aac · ogg · flac · webm**　｜　最大 **100 MB**")
    uploaded = st.file_uploader(
        "拖曳音訊至此，或點擊選擇",
        type=["mp3", "wav", "m4a", "aac", "ogg", "flac", "webm"],
        label_visibility="collapsed",
    )
    if uploaded:
        if st.session_state["_last_uploaded"] != uploaded.name:
            st.session_state["_last_uploaded"] = uploaded.name
            st.session_state["_pending_title"] = Path(uploaded.name).stem
            st.rerun()
        st.audio(uploaded)
        col_info, col_btn = st.columns([5, 1])
        with col_info:
            st.caption(f"檔名：{uploaded.name}　｜　大小：{uploaded.size/1_048_576:.1f} MB")
        with col_btn:
            go = st.button("開始轉錄 →", type="primary", use_container_width=True)
        if go:
            if not groq_key:
                st.error("請在左側輸入 Groq API Key（免費申請：console.groq.com）")
            else:
                with st.spinner("Groq Whisper 轉錄中，請稍候…"):
                    try:
                        audio_data = uploaded.read()
                        segs = transcribe_audio(audio_data, uploaded.name, groq_key, LANG_MAP[language])
                        entries = [
                            {"speaker": speaker_names[0], "text": s["text"],
                             "displayTime": secs_hms(s["start"]), "rawTime": int(s["start"] * 1000)}
                            for s in segs
                        ]
                        info = {
                            "title":        meeting_title,
                            "date":         datetime.now().strftime("%Y/%m/%d %H:%M"),
                            "participants": [speaker_names[0]],
                        }
                        st.session_state.transcript        = entries
                        st.session_state.analysis          = None
                        st.session_state.meeting_info      = info
                        st.session_state.current_record_id = save_to_history(
                            entries, None, info,
                            audio_bytes=audio_data, audio_filename=uploaded.name,
                        )
                        st.success(f"轉錄完成！共 {len(segs)} 段，已自動儲存至歷史記錄。")
                        st.rerun()
                    except Exception as e:
                        st.error(_friendly_error(e))

# ── Tab 2: Record ──────────────────────────────────────────────────────────────
with tab_rec:
    st.info("💡 點擊麥克風錄音，完成後點「轉錄錄音」。需要 Groq API Key。")
    try:
        audio_val = st.audio_input("點擊麥克風開始錄音")
        if audio_val:
            _, col_btn2 = st.columns([5, 1])
            with col_btn2:
                go_rec = st.button("轉錄錄音 →", type="primary", use_container_width=True)
            if go_rec:
                if not groq_key:
                    st.error("請輸入 Groq API Key")
                else:
                    with st.spinner("Groq Whisper 轉錄中…"):
                        try:
                            audio_data = audio_val.read()
                            segs = transcribe_audio(audio_data, "recording.wav", groq_key, LANG_MAP[language])
                            entries = [
                                {"speaker": speaker_names[0], "text": s["text"],
                                 "displayTime": secs_hms(s["start"]), "rawTime": int(s["start"] * 1000)}
                                for s in segs
                            ]
                            info = {
                                "title":        meeting_title,
                                "date":         datetime.now().strftime("%Y/%m/%d %H:%M"),
                                "participants": [speaker_names[0]],
                            }
                            st.session_state.transcript        = entries
                            st.session_state.analysis          = None
                            st.session_state.meeting_info      = info
                            st.session_state.current_record_id = save_to_history(
                                entries, None, info,
                                audio_bytes=audio_data, audio_filename="recording.wav",
                            )
                            st.success("轉錄完成！已自動儲存至歷史記錄。")
                            st.rerun()
                        except Exception as e:
                            st.error(_friendly_error(e))
    except Exception:
        st.warning("即時錄音需要 Streamlit ≥ 1.31，請改用「上傳音檔」分頁。")

# ── Tab 3: 逐字稿 .txt → 會議紀錄 .docx ────────────────────────────────────────
with tab_docx:
    st.markdown(
        "已經有轉好的逐字稿？直接上傳 **.txt**，依「會議行動轉譯官」格式產出 Word 會議紀錄。"
    )
    up_txt = st.file_uploader(
        "上傳逐字稿 .txt", type=["txt"], key="txt_upload", label_visibility="collapsed"
    )

    if up_txt:
        # 換檔時以檔名帶入標題；此處早於下方 widget 建立，可直接寫 session state
        if st.session_state.get("_last_txt_upload") != up_txt.name:
            st.session_state["_last_txt_upload"] = up_txt.name
            st.session_state["docx_title_key"]   = Path(up_txt.name).stem
            st.session_state["docx_bytes"]       = None

        txt_content = _decode_text(up_txt.getvalue())
        raw_entries = parse_transcript_txt(txt_content)
        # Whisper 逐字稿一句一行，先合併成段落再輸出／送 AI
        txt_entries = merge_segments(raw_entries)
        n_spk = len({e["speaker"] for e in txt_entries if e["speaker"]})
        ai_chars = len("\n".join(
            f"{e['speaker']}: {e['text']}" if e["speaker"] and not e.get("cont") else e["text"]
            for e in txt_entries
        ))
        st.caption(
            f"已讀取 {len(raw_entries)} 句 → 合併為 {len(txt_entries)} 段"
            f"　｜　辨識到 {n_spk} 位發言者"
            f"　｜　送 AI 約 {ai_chars:,} 字元"
            + ("（未偵測到「發言者：」格式，將以純段落輸出）" if n_spk == 0 else "")
        )
        if n_spk == 1:
            st.caption(
                "ℹ️ 只偵測到一位發言者——Whisper 不做語者分離，"
                "整份逐字稿會標成同一人。若要區分發言者，請先在 .txt 中改好名稱再上傳。"
            )
        with st.expander("預覽解析結果（前 5 段）"):
            for e in txt_entries[:5]:
                st.markdown(
                    f'<div class="tr-row"><span class="tr-spk">'
                    f'{"" if e.get("cont") else (e["speaker"] or "—")}</span>'
                    f'<span>{e["text"]}</span></div>',
                    unsafe_allow_html=True,
                )

        c_t, c_d = st.columns([3, 2])
        with c_t:
            docx_title = st.text_input("會議標題", key="docx_title_key")
        with c_d:
            docx_date = st.date_input("會議日期", value=datetime.now())

        st.markdown("**格式模板（選填）**")
        st.caption(
            "不上傳 → 採用預設五段結構（Meeting Summary / Topics / Action Items / Notes / "
            "Transcript）。上傳 .txt、.md 或 .docx 範本 → 依範本的欄位與順序填寫。"
        )
        up_tpl = st.file_uploader(
            "上傳格式模板", type=["txt", "md", "docx"], key="tpl_upload",
            label_visibility="collapsed",
        )

        c_o1, c_o2 = st.columns(2)
        with c_o1:
            use_ai = st.checkbox("以 AI 整理摘要／行動事項（需 Groq API Key）", value=True)
        with c_o2:
            keep_tr = st.checkbox("附上完整逐字稿（Transcript）", value=True)
        polish = st.checkbox(
            "並整理逐字稿內容（補標點、標記錯字、分辨發言者角色）",
            value=False, disabled=not (use_ai and keep_tr),
        )
        if polish:
            n_batch = max(1, -(-ai_chars // POLISH_CHUNK_CHARS))
            st.caption(
                f"逐字稿會分 {n_batch} 批送出，受 Groq 免費方案速率限制（12,000 TPM）"
                f"約需 {max(1, n_batch // 3)}–{n_batch // 2 + 1} 分鐘。"
                "任何一批失敗都會保留該批原文，不會少內容。"
            )

        if st.button("產生 Word →", type="primary", key="docx_go"):
            if not txt_entries:
                st.error("這份 .txt 沒有可用內容。")
            elif use_ai and not groq_key:
                st.error("AI 整理需要 Groq API Key，請在左側輸入；或取消勾選只做格式轉檔。")
            else:
                try:
                    tpl_text = read_template_text(up_tpl) if up_tpl else None
                    sections = []
                    if use_ai:
                        if ai_chars > MAX_TRANSCRIPT_CHARS:
                            st.warning(
                                f"⚠️ 逐字稿共 {ai_chars:,} 字元，超過 Groq 免費方案上限"
                                f"（{MAX_TRANSCRIPT_CHARS:,} 字元），摘要將只依首尾片段整理；"
                                "Transcript 區塊仍為完整內容。"
                            )
                        with st.spinner("Llama 整理會議紀錄中…"):
                            sections = generate_minutes(
                                txt_entries,
                                {"title": docx_title, "date": docx_date.strftime("%Y/%m/%d")},
                                groq_key, tpl_text,
                            )["sections"]
                    doc_entries, note = txt_entries, ""
                    if polish and use_ai and keep_tr:
                        pbar = st.progress(0.0, text="整理逐字稿中…")

                        def _tick(i, total, extra=""):
                            pbar.progress(i / total, text=f"整理逐字稿：第 {i}/{total} 批{extra}")

                        try:
                            doc_entries, warns = polish_transcript(
                                txt_entries, groq_key, docx_title, on_progress=_tick
                            )
                        finally:
                            pbar.empty()
                        for w in warns:
                            st.warning(w)
                        note = ("說明：以下逐字稿已由 AI 補上標點、修正明顯的語音辨識錯字"
                                "（標記【原文：…】）並依語意標註發言者角色，內容未經刪減。")

                    st.session_state["docx_bytes"] = build_minutes_docx(
                        docx_title, docx_date.strftime("%Y/%m/%d"),
                        sections, doc_entries, include_transcript=keep_tr,
                        source_name=up_txt.name, transcript_note=note,
                    )
                    st.session_state["docx_name"] = minutes_filename(
                        docx_title, docx_date.strftime("%Y%m%d")
                    )
                    st.success("會議紀錄已產生，可於下方下載。")
                except ModuleNotFoundError:
                    st.error("缺少 python-docx 套件，請執行：pip install python-docx")
                except json.JSONDecodeError:
                    st.error("AI 回傳的內容不是合法 JSON，請再試一次，或取消 AI 整理只做格式轉檔。")
                except Exception as e:
                    st.error(_friendly_error(e))

    if st.session_state.get("docx_bytes"):
        st.download_button(
            "⬇ 下載會議紀錄 (.docx)",
            st.session_state["docx_bytes"],
            st.session_state.get("docx_name", "會議紀錄.docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary", key="docx_dl",
        )

# ── Tab 4: History ─────────────────────────────────────────────────────────────
with tab_hist:
    st.markdown("**載入過去儲存的 JSON 紀錄**")
    loaded_file = st.file_uploader("上傳 .json 檔案", type=["json"],
                                   key="hist_upload", label_visibility="collapsed")
    if loaded_file:
        try:
            rec = json.loads(loaded_file.read().decode("utf-8"))
            if all(k in rec for k in ("transcript", "analysis", "meeting_info")):
                save_to_history(rec["transcript"], rec["analysis"], rec["meeting_info"])
                st.success(f"已載入：{rec['meeting_info'].get('title','')}")
                st.rerun()
            else:
                st.error("格式不正確，請上傳由本系統產生的 .json 檔案")
        except Exception as e:
            st.error(f"載入失敗：{e}")

    st.divider()
    history = st.session_state.history
    if not history:
        st.info("📭 還沒有歷史紀錄。完成 AI 分析後自動儲存，或上傳 .json 檔案。")
    else:
        st.markdown(f"**共 {len(history)} 筆紀錄**")
        col_list, col_detail = st.columns([1, 2])
        with col_list:
            for i, rec in enumerate(history):
                is_active = (st.session_state.hist_idx == i)
                tag = "" if rec.get("analysis") else ' <span style="font-size:.75rem;color:#9ca3af;">（逐字稿）</span>'
                st.markdown(
                    f'<div class="hist-card {"active" if is_active else ""}">'
                    f'<div class="hist-title">{rec["title"]}{tag}</div>'
                    f'<div class="hist-meta">{rec["date"]}</div>'
                    f'<div class="hist-meta">{" · ".join(rec.get("participants", []))}</div></div>',
                    unsafe_allow_html=True,
                )
                c_view, c_del = st.columns(2)
                with c_view:
                    if st.button("查看", key=f"view_{i}", use_container_width=True):
                        st.session_state.hist_idx = i
                        st.rerun()
                with c_del:
                    if st.button("🗑", key=f"del_{i}", use_container_width=True, help="刪除此記錄"):
                        st.session_state.history.pop(i)
                        if st.session_state.hist_idx == i:
                            st.session_state.hist_idx = None
                        elif st.session_state.hist_idx is not None and st.session_state.hist_idx > i:
                            st.session_state.hist_idx -= 1
                        st.rerun()
        with col_detail:
            idx = st.session_state.hist_idx
            if idx is not None and idx < len(history):
                rec = history[idx]
                ab = rec.get("audio_bytes")
                af = rec.get("audio_filename", "recording")
                if rec["analysis"] is None:
                    st.info("📝 此記錄尚未進行 AI 分析，僅顯示逐字稿。")
                    info = rec["meeting_info"]
                    st.markdown(f"**{info.get('title','')}**　{info.get('date','')}")
                    if ab:
                        st.audio(ab)
                        st.download_button("⬇ 下載原始錄音", ab, af,
                                           key=f"hist_{idx}_dl_audio_raw")
                    fname = info.get("title", "transcript").replace(" ", "_")
                    st.download_button(
                        "⬇ 下載逐字稿",
                        "\n".join(f"{e['speaker']}: {e['text']}" for e in rec["transcript"]),
                        f"{fname}.txt", key=f"hist_{idx}_dl_txt_raw",
                    )
                    for e in rec["transcript"]:
                        st.markdown(
                            f'<div class="tr-row"><span class="tr-spk">{e["speaker"]}:</span>'
                            f'<span>{e["text"]}</span></div>',
                            unsafe_allow_html=True,
                        )
                else:
                    render_results(rec["analysis"], rec["meeting_info"], rec["transcript"],
                                   key_prefix=f"hist_{idx}", audio_bytes=ab, audio_filename=af)
            else:
                st.info("← 點擊左側紀錄查看內容")

# ── Transcript editor ──────────────────────────────────────────────────────────
if st.session_state.transcript:
    st.divider()
    with st.expander(
        f"📝 逐字稿編輯　（{len(st.session_state.transcript)} 段，可修改文字 / 指定發言者）",
        expanded=True,
    ):
        for i, entry in enumerate(st.session_state.transcript):
            c_sp, c_time, c_txt = st.columns([2, 1, 7])
            with c_sp:
                idx = speaker_names.index(entry["speaker"]) if entry["speaker"] in speaker_names else 0
                st.session_state.transcript[i]["speaker"] = st.selectbox(
                    "sp", speaker_names, index=idx, key=f"sel{i}", label_visibility="collapsed")
            with c_time:
                st.markdown(
                    f"<p style='color:#9ca3af;font-size:.8rem;padding-top:.55rem'>{entry['displayTime']}</p>",
                    unsafe_allow_html=True)
            with c_txt:
                st.session_state.transcript[i]["text"] = st.text_input(
                    "txt", value=entry["text"], key=f"txt{i}", label_visibility="collapsed")

    c_ana, c_dl, _ = st.columns([2, 2, 4])
    with c_ana:
        do_analyze = st.button("🤖 AI 分析", type="primary", use_container_width=True)
    with c_dl:
        st.download_button(
            "⬇ 下載逐字稿",
            "\n".join(f"{e['speaker']}: {e['text']}" for e in st.session_state.transcript),
            f"transcript_{datetime.now():%Y%m%d_%H%M}.txt",
            use_container_width=True,
        )

    if do_analyze:
        if not groq_key:
            st.error("請輸入 Groq API Key")
        else:
            full_txt = "\n".join(e["text"] for e in st.session_state.transcript)
            if len(full_txt) > MAX_TRANSCRIPT_CHARS:
                st.warning(
                    f"⚠️ 逐字稿共 {len(full_txt):,} 字元，超過 Groq 免費方案上限（{MAX_TRANSCRIPT_CHARS:,} 字元）。"
                    " 系統將自動保留首尾最重要的片段進行分析。如需完整分析，請升級至 Groq Dev Tier。"
                )
            with st.spinner("Llama 分析中，請稍候…"):
                try:
                    info = {
                        "title":        meeting_title,
                        "date":         datetime.now().strftime("%Y/%m/%d %H:%M"),
                        "participants": list(dict.fromkeys(e["speaker"] for e in st.session_state.transcript)),
                    }
                    result = analyze_with_groq(st.session_state.transcript, info, groq_key)
                    st.session_state.analysis          = result
                    st.session_state.meeting_info      = info
                    st.session_state.current_record_id = save_to_history(
                        st.session_state.transcript, result, info,
                        st.session_state.current_record_id,
                    )
                    st.rerun()
                except Exception as e:
                    st.error(_friendly_error(e))

# ── Current results ────────────────────────────────────────────────────────────
if st.session_state.analysis:
    st.divider()
    st.markdown('<div class="sec-title">📋 本次分析結果</div>', unsafe_allow_html=True)
    render_results(
        st.session_state.analysis,
        st.session_state.meeting_info,
        st.session_state.transcript,
    )

# ── 每次 run 結束時將歷史記錄同步至 localStorage ───────────────────────────────
# 必須等 _history_loaded=True 才能寫入，否則第一次 render JS 尚未執行（回傳 0），
# history 還是空的，會把 localStorage 的舊資料覆蓋掉。
if st.session_state["_history_loaded"]:
    _persist_history()
